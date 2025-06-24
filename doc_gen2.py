import streamlit as st
import google.generativeai as genai
import os
from docx import Document
import io
from datetime import datetime
import PyPDF2
import pytesseract
from PIL import Image
import json

# -------------------- CUSTOM CSS STYLING --------------------

def apply_custom_css():
    st.markdown("""
    <style>
    /* Sidebar headers */
    .sidebar .element-container h2, .sidebar .element-container h3 {
        color: #e6edf3 !important;
    }
    
    /* Sidebar text */
    .sidebar .element-container p, .sidebar .element-container li {
        color: #8b949e !important;
    }
    
    /* File uploader */
    .stFileUploader > div > div > div > div {
        color: #8b949e !important;
    }
    
    .stFileUploader > div > div > div > button {
        color: #e6edf3 !important;
    }
    
    /* Selectbox */
    .stSelectbox > div > div > div {
        color: #e6edf3 !important;
    }
    
    /* Main title */
    .main-title {
        color: #e6edf3 !important;
    }
    
    /* Subtitle */
    .main-subtitle {
        color: #8b949e !important;
    }
    
    /* Section headers */
    .section-header-patient {
        color: rgb(0, 212, 170) !important;
    }
    
    .section-header-claim {
        color: rgb(0, 212, 170) !important;
    }
    
    /* Input labels */
    .stTextInput > label, .stTextArea > label, .stDateInput > label, .stNumberInput > label {
        color: #8b949e !important;
    }
    
    /* Input fields */
    .stTextInput > div > div > input {
        color: #e6edf3 !important;
    }
    
    .stTextArea > div > div > textarea {
        color: #e6edf3 !important;
    }
    
    .stDateInput > div > div > input {
        color: #e6edf3 !important;
    }
    
    .stNumberInput > div > div > input {
        color: #e6edf3 !important;
    }
    
    /* Placeholder text */
    .stTextInput input::placeholder,
    .stTextArea textarea::placeholder {
        color: #8b949e !important;
    }
    
    /* Number input buttons */
    .stNumberInput button {
        background-color: rgb(0, 212, 170) !important;
        color: #f0f6fc !important;
    }
    
    /* Generate Document button */
    .stButton > button[kind="primary"] {
        background-color: rgb(0, 212, 170) !important;
        color: #f0f6fc !important;
    }
    
    .stButton > button[kind="primary"]:hover {
        background-color: #00a085 !important;
        color: #f0f6fc !important;
    }
    
    /* Reset button styling */
    .stButton > button:not([kind="primary"]) {
        color: #d29922 !important;
    }
    
    /* Sidebar instructions */
    .sidebar-instructions {
        color: #e6edf3 !important;
    }
    
    .sidebar-instructions-text {
        color: #8b949e !important;
    }
    
    /* Download button */
    .stDownloadButton > button {
        background-color: rgb(0, 212, 170) !important;
        color: #f0f6fc !important;
    }
    
    .stDownloadButton > button:hover {
        background-color: #00a085 !important;
    }
    
    /* Text area for generated content */
    .stTextArea[data-testid="stTextArea"] textarea {
        color: #e6edf3 !important;
    }
    
    /* Footer styling */
    .footer-text {
        color: #8b949e !important;
        text-align: center !important;
    }
    </style>
    """, unsafe_allow_html=True)

# -------------------- CONFIGURATION --------------------

st.set_page_config(
    page_title="AI Healthcare Document Generator",
    page_icon="📄",
    layout="wide"
)

# Apply custom CSS
apply_custom_css()

@st.cache_resource
def init_gemini_client():
    try:
        api_key = st.secrets.get("GEMINI_API_KEY") or os.getenv("GEMINI_API_KEY")
        if not api_key:
            st.error("Gemini API key not found. Please set it in Streamlit secrets or environment variables.")
            st.stop()
        genai.configure(api_key=api_key)
        return genai.GenerativeModel('gemini-1.5-flash')
    except Exception as e:
        st.error(f"Error initializing Gemini client: {str(e)}")
        return None

# -------------------- DOCUMENT TEXT EXTRACTION --------------------

def extract_text_from_pdf(pdf_file):
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {str(e)}")
        return None

def extract_text_from_image(image_file):
    try:
        image = Image.open(image_file)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        st.error(f"Error extracting text from image: {str(e)}")
        return None

# -------------------- AI INFORMATION EXTRACTION --------------------

def clean_json_response(response_text):
    """Removes markdown and extracts JSON object from Gemini response."""
    response_text = response_text.strip()
    # Remove markdown code block markers if present
    if response_text.startswith("```json"):
        response_text = response_text[7:]
    if response_text.startswith("```"):
        response_text = response_text[3:]
    if response_text.endswith("```"):
        response_text = response_text[:-3]
    response_text = response_text.strip()
    # Extract JSON object
    start = response_text.find('{')
    end = response_text.rfind('}') + 1
    if start == -1 or end == 0:
        return None
    return response_text[start:end]

def extract_information_from_document(document_text, document_type):
    model = init_gemini_client()
    if not model:
        return None

    prompt = f"""
    Extract relevant information from the following {document_type} document text for healthcare form filling.

    Document Text:
    {document_text[:2000]}

    IMPORTANT: You must return ONLY valid JSON format. Do not include any explanatory text before or after the JSON.

    Extract and return this exact JSON structure:
    {{
        "patient_name": "Full name of the patient or Not found",
        "policy_number": "Insurance policy or member ID or Not found",
        "date_of_birth": "Date of birth in YYYY-MM-DD format or Not found",
        "phone": "Phone number or Not found",
        "email": "Email address or Not found", 
        "address": "Full address or Not found",
        "diagnosis": "Medical diagnosis or condition or Not found",
        "treatment": "Treatment or procedure details or Not found",
        "service_date": "Date of service in YYYY-MM-DD format or Not found",
        "provider_name": "Healthcare provider name or Not found",
        "claim_amount": "Claim amount numbers only or Not found",
        "insurance_company": "Insurance company name or Not found"
    }}

    Return only the JSON object, nothing else.
    """

    try:
        generation_config = genai.types.GenerationConfig(
            temperature=0.1,
            max_output_tokens=800,
            top_p=0.8,
            top_k=40
        )
        response = model.generate_content(
            prompt,
            generation_config=generation_config
        )
        if not response or not response.text:
            st.error("Empty response from Gemini API")
            return None

        response_text = clean_json_response(response.text)
        if not response_text:
            st.error("No valid JSON found in response")
            return None
        extracted_data = json.loads(response_text)
        return extracted_data

    except json.JSONDecodeError as e:
        st.error(f"JSON parsing error: {str(e)}")
        return None
    except Exception as e:
        st.error(f"Error extracting information: {str(e)}")
        return None

def extract_information_simple(document_text, document_type):
    model = init_gemini_client()
    if not model:
        return None

    prompt = f"""
    Extract information from this {document_type} document and provide simple answers:

    Document Text:
    {document_text[:1500]}

    Please answer these questions based on the document:
    1. Patient name:
    2. Policy number:
    3. Date of birth:
    4. Phone number:
    5. Email:
    6. Address:
    7. Diagnosis:
    8. Treatment:
    9. Service date:
    10. Claim amount:

    Answer "Not found" if information is not available.
    """

    try:
        response = model.generate_content(prompt)
        if response and response.text:
            lines = response.text.strip().split('\n')
            extracted_data = {}
            for line in lines:
                if ':' in line:
                    key, value = line.split(':', 1)
                    key = key.strip().lower().replace(' ', '_')
                    value = value.strip()
                    if 'name' in key:
                        extracted_data['patient_name'] = value
                    elif 'policy' in key:
                        extracted_data['policy_number'] = value
                    elif 'birth' in key:
                        extracted_data['date_of_birth'] = value
                    elif 'phone' in key:
                        extracted_data['phone'] = value
                    elif 'email' in key:
                        extracted_data['email'] = value
                    elif 'address' in key:
                        extracted_data['address'] = value
                    elif 'diagnosis' in key:
                        extracted_data['diagnosis'] = value
                    elif 'treatment' in key:
                        extracted_data['treatment'] = value
                    elif 'service' in key:
                        extracted_data['service_date'] = value
                    elif 'amount' in key:
                        extracted_data['claim_amount'] = value
            return extracted_data
    except Exception as e:
        st.error(f"Error in simple extraction: {str(e)}")
        return None

def auto_fill_form(extracted_data):
    if not extracted_data:
        return
    st.session_state.extracted_data = extracted_data
    st.success("✅ Document processed! Form fields will be auto-filled below.")

# -------------------- DOCUMENT GENERATION --------------------

def generate_document_content(document_type, patient_data, claim_details):
    model = init_gemini_client()
    if not model:
        return None

    prompt = f"""
    Generate a professional {document_type} based on the following information:

    Patient Information:
    - Name: {patient_data.get('name', '')}
    - Policy Number: {patient_data.get('policy_number', '')}
    - Date of Birth: {patient_data.get('dob', '')}
    - Contact Information: {patient_data.get('contact', '')}

    Claim Details:
    - Service Date: {claim_details.get('service_date', '')}
    - Diagnosis: {claim_details.get('diagnosis', '')}
    - Treatment: {claim_details.get('treatment', '')}
    - Claim Amount: {claim_details.get('amount', '')}
    - Reason for Claim/Appeal: {claim_details.get('reason', '')}

    Please generate a formal, professional document that includes:
    1. Proper business letter formatting with recipient address placeholder
    2. Clear statement of the claim/appeal
    3. Supporting medical information and justification
    4. Specific requested action
    5. Professional closing with signature line

    Make it persuasive yet factual, following healthcare industry standards.
    The tone should be professional and respectful.
    Use Indian Rupees (₹) for all monetary amounts.
    Format should be suitable for Indian healthcare system.
    """

    try:
        generation_config = genai.types.GenerationConfig(
            temperature=0.3,
            max_output_tokens=1500,
            top_p=0.8,
            top_k=40
        )
        response = model.generate_content(
            prompt,
            generation_config=generation_config
        )
        return response.text
    except Exception as e:
        st.error(f"Error generating content: {str(e)}")
        return None

def create_word_document(content, doc_type, patient_name):
    doc = Document()
    # Add header
    sections = doc.sections
    header = sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = f"{doc_type} - {patient_name}"
    # Add title
    title = doc.add_heading(doc_type, 0)
    title.alignment = 1  # Center alignment
    # Add date
    date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = 2  # Right alignment
    # Add content
    doc.add_paragraph(content)
    # Add footer
    footer = sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Generated by AI Healthcare Document Generator"
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

def generate_and_display_document(doc_type, name, policy, dob, contact, 
                                service_date, diagnosis, treatment, amount, reason):
    patient_data = {
        'name': name,
        'policy_number': policy,
        'dob': str(dob),
        'contact': contact
    }
    claim_details = {
        'service_date': str(service_date),
        'diagnosis': diagnosis,
        'treatment': treatment,
        'amount': f"₹{amount:,.2f}" if amount > 0 else "Not specified",
        'reason': reason
    }
    with st.spinner("🤖 Generating document with Gemini AI..."):
        content = generate_document_content(doc_type, patient_data, claim_details)
        if content:
            st.success("✅ Document generated successfully!")
            st.subheader("📄 Generated Document")
            st.text_area("Document Content", content, height=400, key="generated_content")
            doc_file = create_word_document(content, doc_type, name)
            st.download_button(
                label="📥 Download as Word Document",
                data=doc_file,
                file_name=f"{doc_type.replace(' ', '_')}_{name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.subheader("✏️ Edit and Regenerate")
            if st.button("🔄 Generate New Version"):
                st.rerun()

# -------------------- MAIN APP --------------------

def main():
    # Main title with custom styling
    st.markdown('<h1 class="main-title">🏥 AI Healthcare Document Generator</h1>', unsafe_allow_html=True)
    st.markdown('<p class="main-subtitle">Generate professional claim letters, appeals, and healthcare documents automatically using Google Gemini AI</p>', unsafe_allow_html=True)

    if 'extracted_data' not in st.session_state:
        st.session_state.extracted_data = {}

    # Sidebar for settings and document upload
    with st.sidebar:
        st.markdown('<h2 style="color: #e6edf3;">📁 Document Upload & Auto-Fill</h2>', unsafe_allow_html=True)
        uploaded_file = st.file_uploader(
            "Upload Document for Auto-Fill",
            type=['pdf', 'png', 'jpg', 'jpeg'],
            help="Upload medical records, insurance cards, or other relevant documents"
        )
        if uploaded_file is not None:
            document_type = st.selectbox(
                "Document Type",
                ["Medical Record", "Insurance Card", "Previous Claim", "Medical Bill", "Other"]
            )
            if st.button("🔍 Extract Information", type="primary"):
                with st.spinner("Processing document..."):
                    if uploaded_file.type == "application/pdf":
                        document_text = extract_text_from_pdf(uploaded_file)
                    else:
                        document_text = extract_text_from_image(uploaded_file)
                    if document_text:
                        extracted_info = extract_information_from_document(document_text, document_type)
                        if not extracted_info:
                            st.warning("JSON extraction failed, trying simple extraction...")
                            extracted_info = extract_information_simple(document_text, document_type)
                        if extracted_info:
                            auto_fill_form(extracted_info)
                            st.subheader("📋 Extracted Information")
                            for key, value in extracted_info.items():
                                if value and value != "Not found":
                                    st.write(f"**{key.replace('_', ' ').title()}:** {value}")
                        else:
                            st.error("Could not extract information from the document")
        
        st.markdown("---")
        st.markdown('<h2 style="color: #e6edf3;">⚙️ Document Settings</h2>', unsafe_allow_html=True)
        document_type = st.selectbox(
            "Select Document Type",
            [
                "Insurance Claim Letter", 
                "Appeal Letter", 
                "Prior Authorization Request", 
                "Reimbursement Claim",
                "Medical Necessity Letter",
                "Coverage Determination Appeal"
            ]
        )
        
        st.markdown("---")
        st.markdown('<h3 class="sidebar-instructions">📋 Instructions</h3>', unsafe_allow_html=True)
        st.markdown("""
        <div class="sidebar-instructions-text">
        1. <strong>Upload a document</strong> (optional) to auto-fill form fields<br>
        2. Fill in remaining required fields (marked with *)<br>
        3. Select the appropriate document type<br>
        4. Click 'Generate Document'<br>
        5. Review and download the generated document
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("---")
        if st.button("🔄 Reset All Fields", key="reset_btn"):
            st.session_state.extracted_data = {}
            st.rerun()

    # Main form with auto-fill capability
    col1, col2 = st.columns(2)
    extracted = st.session_state.extracted_data

    with col1:
        st.markdown('<h3 class="section-header-patient">👤 Patient Information</h3>', unsafe_allow_html=True)
        patient_name = st.text_input(
            "Patient Full Name*", 
            value=extracted.get('patient_name', '') if extracted.get('patient_name') != 'Not found' else '',
            placeholder="Rajesh Kumar"
        )
        policy_number = st.text_input(
            "Policy/Member ID*", 
            value=extracted.get('policy_number', '') if extracted.get('policy_number') != 'Not found' else '',
            placeholder="HDFC123456789"
        )
        dob_value = datetime(1990, 1, 1)
        if extracted.get('date_of_birth') and extracted.get('date_of_birth') != 'Not found':
            try:
                dob_value = datetime.strptime(extracted.get('date_of_birth'), '%Y-%m-%d')
            except ValueError:
                pass
        dob = st.date_input("Date of Birth", value=dob_value)
        contact_value = ""
        if extracted.get('phone') and extracted.get('phone') != 'Not found':
            contact_value += f"Phone: {extracted.get('phone')}\n"
        if extracted.get('email') and extracted.get('email') != 'Not found':
            contact_value += f"Email: {extracted.get('email')}\n"
        if extracted.get('address') and extracted.get('address') != 'Not found':
            contact_value += f"Address: {extracted.get('address')}"
        contact_info = st.text_area(
            "Contact Information", 
            value=contact_value,
            placeholder="Phone: +91 98765 43210\nEmail: patient@email.com\nAddress: 123 MG Road, Mumbai, Maharashtra 400001"
        )

    with col2:
        st.markdown('<h3 class="section-header-claim">🏥 Claim Details</h3>', unsafe_allow_html=True)
        service_date_value = datetime.now()
        if extracted.get('service_date') and extracted.get('service_date') != 'Not found':
            try:
                service_date_value = datetime.strptime(extracted.get('service_date'), '%Y-%m-%d')
            except ValueError:
                pass
        service_date = st.date_input("Service Date", value=service_date_value)
        diagnosis = st.text_area(
            "Diagnosis/Condition*", 
            value=extracted.get('diagnosis', '') if extracted.get('diagnosis') != 'Not found' else '',
            placeholder="Primary diagnosis code and description"
        )
        treatment = st.text_area(
            "Treatment/Service*", 
            value=extracted.get('treatment', '') if extracted.get('treatment') != 'Not found' else '',
            placeholder="Detailed description of treatment or service provided"
        )
        claim_amount_value = 0.0
        if extracted.get('claim_amount') and extracted.get('claim_amount') != 'Not found':
            try:
                amount_str = ''.join(filter(str.isdigit, str(extracted.get('claim_amount'))))
                if amount_str:
                    claim_amount_value = float(amount_str)
            except ValueError:
                pass
        claim_amount = st.number_input(
            "Claim Amount (₹)", 
            min_value=0.0, 
            format="%.2f",
            value=claim_amount_value,
            help="Enter 0 if amount is not applicable"
        )
        reason = st.text_area(
            "Reason for Claim/Appeal*", 
            placeholder="Detailed explanation of why this claim should be approved or why the appeal should be considered"
        )

    st.markdown("---")
    
    # Custom styled generate button
    generate_clicked = st.button("📄 Generate Document", type="primary", use_container_width=True)
    
    if generate_clicked:
        required_fields = [patient_name, policy_number, diagnosis, treatment, reason]
        if not all(field.strip() for field in required_fields):
            st.error("❌ Please fill in all required fields marked with *")
        else:
            generate_and_display_document(
                document_type, patient_name, policy_number, 
                dob, contact_info, service_date, diagnosis, 
                treatment, claim_amount, reason
            )

    st.markdown("---")
    st.markdown(
        """
        <div class="footer-text">
            AI Healthcare Document Generator | Powered by Google Gemini & Streamlit
        </div>
        """, 
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()